# -*- coding: utf-8 -*-
"""Guide RDV commercial ANDRAGOPS x COUSIN GROUP — offre réelle : SST, MAC SST,
gestes & postures, PRAP IBC, incendie, cohésion d'entreprise."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU = RGBColor(0x1F, 0x3A, 0x5F)
ROUGE = RGBColor(0xB3, 0x1B, 0x1B)
GRIS = RGBColor(0x44, 0x44, 0x44)
VERT = RGBColor(0x1E, 0x6B, 0x3A)


def set_base_style(doc):
    s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(11); s.font.color.rgb = GRIS


def add_cover(doc, title, subtitle, meta_lines):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = BLEU
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(15); r.font.color.rgb = ROUGE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—' * 20); r.font.color.rgb = BLEU
    for line in meta_lines:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(11); r.font.color.rgb = GRIS
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = BLEU
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '4'), ('w:color', 'B31B1B')):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12.5); r.font.bold = True; r.font.color.rgb = ROUGE
    return p


def para(doc, text, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.bold = bold
    return p


def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    for i, part in enumerate(text.split('**')):
        r = p.add_run(part)
        if i % 2 == 1:
            r.font.bold = True
    return p


def script_line(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + ' '); r.font.bold = True; r.font.color.rgb = ROUGE
    r2 = p.add_run('« ' + text + ' »'); r2.font.italic = True
    return p


def add_table(doc, header, rows, widths):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(header):
        c = table.rows[0].cells[i]; c.width = Inches(widths[i])
        rr = c.paragraphs[0].add_run(htxt); rr.font.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = BLEU
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            rr = cells[i].paragraphs[0].add_run(val); rr.font.size = Pt(9.5)
    doc.add_paragraph()
    return table


def add_objection(doc, objection, reponse):
    p = doc.add_paragraph(); r = p.add_run('« ' + objection + ' »'); r.font.bold = True; r.font.color.rgb = BLEU
    p2 = doc.add_paragraph(); r2 = p2.add_run('→ Réponse : '); r2.font.bold = True; r2.font.color.rgb = VERT
    p2.add_run(reponse); p2.paragraph_format.space_after = Pt(10)


def fill_lines(doc, n=3):
    for _ in range(n):
        doc.add_paragraph("……………………………………………………………………………………………………")


doc = Document()
set_base_style(doc)

add_cover(
    doc,
    "Guide de rendez-vous commercial",
    "ANDRAGOPS Académie  ✕  COUSIN GROUP",
    [
        "Échange téléphonique avec la Responsable RH",
        "Formations prévention & santé-sécurité au travail + cohésion d'équipe",
        "Angle : sur-mesure & ancrage local (Hauts-de-France)",
        "Contact : +33 (0)6 35 25 89 22 — Wervicq-Sud (59)",
        "",
        "Guide préparé le 19 juin 2026",
    ],
)

# 1. Objectif
h1(doc, "1. Objectif & état d'esprit du rendez-vous")
para(doc, "Échange de découverte : présenter brièvement ANDRAGOPS, faire parler la RH sur leurs enjeux "
          "sécurité / prévention / cohésion, puis proposer une piste sur-mesure. La sécurité au travail "
          "est en grande partie une obligation réglementaire : c'est un levier de vente puissant.")
add_table(doc, ["Élément", "Repère"], [
    ["But de l'appel", "Qualifier les besoins SST/prévention, décrocher un RDV de cadrage"],
    ["Ce que je vends", "Formations sécurité obligatoires/recommandées + cohésion, en sur-mesure et local"],
    ["Posture", "Expert-conseil prévention (80 % écoute / 20 % parole)"],
    ["Résultat visé", "Un audit de leurs besoins sécurité + une proposition chiffrée"],
], [2.0, 4.3])

# 2. Catalogue ANDRAGOPS
h1(doc, "2. Votre catalogue ANDRAGOPS (à présenter clairement)")
add_table(doc, ["Formation", "À quoi ça sert", "Récurrence / cible"], [
    ["SST — Sauveteur Secouriste du Travail",
     "Former des salariés à porter secours et à prévenir les risques",
     "Initiale (~14 h). Cible : ouvriers atelier, encadrants"],
    ["MAC SST (recyclage)",
     "Maintien & actualisation des compétences SST",
     "Obligatoire tous les 24 mois → revenu récurrent"],
    ["Gestes & postures",
     "Adopter les bons gestes pour éviter les blessures",
     "Tout salarié manipulant des charges"],
    ["PRAP IBC",
     "Prévention des risques liés à l'activité physique (Industrie/BTP/Commerce)",
     "Postes physiques, port de charges, TMS"],
    ["Incendie",
     "Manipulation extincteurs, évacuation, équipiers 1ʳᵉ intervention",
     "Tout le site (exercices recommandés ~tous les 6 mois)"],
    ["Cohésion d'entreprise",
     "Activités de team building, renforcer le collectif",
     "Équipes, managers, séminaires"],
], [1.7, 2.5, 2.1])
para(doc, "→ Ajustez les durées/tarifs avec vos chiffres réels. Notez vos tarifs indicatifs ci-dessous :")
fill_lines(doc, 2)

# 3. Pourquoi Cousin a besoin de vous
h1(doc, "3. Pourquoi le COUSIN GROUP est un prospect idéal")
para(doc, "Site industriel ~150 personnes (cordages, composites, médical) : manipulation de fibres et "
          "bobines, machines de tressage/câblage, port de charges, salle blanche côté médical. "
          "Beaucoup de besoins sécurité concrets et souvent réglementaires.")
bullet(doc, "**SST** : il faut des secouristes formés dans chaque atelier (recommandation INRS/Assurance Maladie : "
            "au moins un SST par équipe / zone de travail).")
bullet(doc, "**MAC SST** : tous leurs SST déjà formés doivent être recyclés tous les 2 ans → besoin récurrent.")
bullet(doc, "**Gestes & postures / PRAP IBC** : les TMS (troubles musculo-squelettiques) sont la 1ʳᵉ cause "
            "de maladies professionnelles dans l'industrie → enjeu santé + absentéisme + coût.")
bullet(doc, "**Incendie** : obligation d'organiser des exercices d'évacuation et de former le personnel à la "
            "manipulation des extincteurs (Code du travail) ; matières et machines = risque réel.")
bullet(doc, "**Cohésion** : une maison familiale de 175 ans, en croissance et qui recrute → fédérer "
            "anciens et nouveaux, renforcer l'esprit d'équipe.")

# 4. Pitch
h1(doc, "4. Pitch ANDRAGOPS (30–60 secondes)")
script_line(doc, "Accroche :", "Bonjour, [Prénom Nom], d'ANDRAGOPS Académie. Merci de me consacrer "
            "quelques minutes. On accompagne les entreprises industrielles sur la sécurité et la prévention "
            "des risques au travail.")
script_line(doc, "Qui :", "Concrètement, on forme vos équipes au secourisme (SST), aux gestes et postures, "
            "à la PRAP, à l'incendie — et on propose aussi des activités de cohésion. Le tout en sur-mesure "
            "et en proximité, ici dans les Hauts-de-France.")
script_line(doc, "Pourquoi vous :", "Avec un site comme le vôtre, il y a forcément des obligations à tenir "
            "(secouristes, évacuation, prévention des TMS) — et c'est exactement notre métier.")
script_line(doc, "Transition :", "Avant de vous en dire plus, j'aimerais comprendre où vous en êtes "
            "aujourd'hui sur ces sujets. Je peux vous poser quelques questions ?")
para(doc, "Votre version personnalisée :"); fill_lines(doc, 3)

# 5. Questions de découverte
h1(doc, "5. Questions de découverte (faire parler la RH)")
h2(doc, "Secourisme / SST")
bullet(doc, "Combien de Sauveteurs Secouristes du Travail avez-vous aujourd'hui, et dans quels ateliers ?")
bullet(doc, "Vos SST sont-ils à jour de leur recyclage (MAC tous les 24 mois) ? Qui suit les échéances ?")
h2(doc, "Risques physiques / TMS")
bullet(doc, "Avez-vous des postes avec port de charges ou gestes répétitifs ? Des arrêts liés au dos / TMS ?")
bullet(doc, "Vos équipes ont-elles déjà été formées aux gestes et postures ou à la PRAP ?")
h2(doc, "Incendie")
bullet(doc, "Quand a eu lieu votre dernier exercice d'évacuation ? Avez-vous des équipiers de première intervention formés ?")
bullet(doc, "Votre personnel sait-il manipuler un extincteur ?")
h2(doc, "Cohésion & contexte RH")
bullet(doc, "Vous recrutez en ce moment : comment se passe l'intégration et la cohésion des équipes ?")
bullet(doc, "Comment gérez-vous votre plan de formation sécurité aujourd'hui (interne, prestataires) ?")
h2(doc, "Décision / budget")
bullet(doc, "Qui décide des formations sécurité ? Avez-vous un budget mobilisable via votre OPCO cette année ?")

# 6. Arguments
h1(doc, "6. Vos arguments différenciants")
bullet(doc, "**Obligation = priorité** : la sécurité n'est pas une option, c'est la loi — vous les aidez à être en règle.")
bullet(doc, "**Sur-mesure** : formations bâties sur LEURS postes réels (atelier de tressage, manutention de bobines, salle blanche).")
bullet(doc, "**Local & réactif** : intervention sur site à Wervicq-Sud, planning calé sur la production, interlocuteur unique.")
bullet(doc, "**Approche globale** : un seul partenaire pour SST, gestes & postures, PRAP, incendie ET cohésion.")
bullet(doc, "**Récurrence maîtrisée** : suivi des recyclages MAC SST → vous ne ratez jamais une échéance.")
bullet(doc, "**Financement** : [si Qualiopi] prise en charge possible via l'OPCO et le plan de développement des compétences.")

# 7. Objections
h1(doc, "7. Objections fréquentes & réponses")
add_objection(doc, "On a déjà un organisme pour la sécurité.",
              "« Parfait, beaucoup de nos clients en ont un. On intervient souvent en complément, sur ce qui "
              "manque ou pour les recyclages. On peut commencer par un seul besoin, par exemple le MAC SST. »")
add_objection(doc, "Nos salariés sont déjà formés.",
              "« Très bien. Et leurs recyclages sont-ils à jour ? Le SST se recycle tous les 2 ans. Je peux vous "
              "faire un point gratuit sur vos échéances. »")
add_objection(doc, "On n'a pas le temps, on est en production.",
              "« C'est pour ça qu'on vient sur site, en formats courts adaptés à vos horaires d'atelier, "
              "sans casser la production. »")
add_objection(doc, "C'est un budget en plus.",
              "« La plupart de ces formations sont finançables par votre OPCO. Et une formation prévention "
              "coûte bien moins cher qu'un accident, un arrêt de travail ou un contrôle. »")
add_objection(doc, "Envoyez-moi une plaquette.",
              "« Je vous l'envoie. Mais le plus utile serait 30 min pour faire le point sur vos obligations "
              "réelles, et je reviens avec une proposition ciblée. On se cale un créneau ? »")

# 8. Financement
h1(doc, "8. Financement (bon à savoir)")
bullet(doc, "**Qualiopi** : à mentionner si vous l'avez — c'est ce qui ouvre les financements OPCO.")
bullet(doc, "**Habilitation INRS** : pour SST/PRAP, valorisez votre habilitation (gage de sérieux) si vous l'avez.")
bullet(doc, "**OPCO** : Cousin (industrie textile/plasturgie) dépend a priori de l'**OPCO 2i** — à confirmer.")
bullet(doc, "**Plan de développement des compétences** : le canal de financement côté entreprise.")

# 9. Closing
h1(doc, "9. Conclure : la prochaine étape")
script_line(doc, "Proposition :", "Ce que je vous propose : un point de 30 min, sur site ou en visio, pour "
            "faire l'état de vos obligations (SST, évacuation, TMS) et vos échéances de recyclage. Je reviens "
            "avec un plan et un devis sur-mesure.")
bullet(doc, "Proposez **2 créneaux précis**.")
bullet(doc, "Validez : interlocuteur, email, prochaine étape, ce que vous envoyez d'ici là.")
bullet(doc, "Reformulez les besoins entendus avant de raccrocher.")

# 10. Check-list + notes
h1(doc, "10. Check-list avant l'appel")
bullet(doc, "Dossier de recherche COUSIN GROUP + ce guide sous les yeux.")
bullet(doc, "Mes tarifs/durées par formation + 1 exemple de client industriel satisfait.")
bullet(doc, "Mes créneaux de RDV disponibles.")
bullet(doc, "Endroit calme, bon réseau, de quoi noter, sourire (s'entend au téléphone).")

h1(doc, "11. Notes de l'appel")
para(doc, "Interlocuteur / fonction :"); fill_lines(doc, 1)
para(doc, "Nb de SST actuels / échéances recyclage :"); fill_lines(doc, 1)
para(doc, "Besoins exprimés (gestes/PRAP, incendie, cohésion) :"); fill_lines(doc, 3)
para(doc, "Budget / OPCO / décideurs :"); fill_lines(doc, 1)
para(doc, "Prochaine étape & date :"); fill_lines(doc, 2)

doc.save('/home/user/code-rouge/Guide_RDV_Commercial_COUSIN_GROUP.docx')
print("Guide commercial (offre réelle) OK")
